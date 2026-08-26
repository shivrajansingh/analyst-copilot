"""Multi-format parsing: detection, per-format segmentation, and Markdown storage."""

import csv

import pytest

from analyst_copilot.parsing import (
    DocumentFormat,
    MarkdownPageStore,
    SegmentKind,
    detect_format,
    parse_document,
)
from analyst_copilot.parsing.formats import UnsupportedFormat
from analyst_copilot.parsing.markdown import render_table


# -- format detection ----------------------------------------------------- #

def test_detects_each_supported_extension(tmp_path):
    cases = {
        "a.htm": DocumentFormat.HTML,
        "a.html": DocumentFormat.HTML,
        "a.docx": DocumentFormat.DOCX,
        "a.xlsx": DocumentFormat.XLSX,
        "a.csv": DocumentFormat.CSV,
        "a.md": DocumentFormat.MARKDOWN,
        "a.txt": DocumentFormat.TEXT,
    }
    for name, expected in cases.items():
        path = tmp_path / name
        path.write_bytes(b"x")
        assert detect_format(path) is expected, name


def test_pdf_is_detected_from_its_magic_bytes_whatever_it_is_called(tmp_path):
    """An analyst who saves a PDF as .txt should still get the PDF parser."""
    path = tmp_path / "statement.txt"
    path.write_bytes(b"%PDF-1.7\n%\xe2\xe3\xcf\xd3\n")
    assert detect_format(path) is DocumentFormat.PDF


def test_a_file_named_pdf_that_is_not_one_is_rejected(tmp_path):
    path = tmp_path / "notes.pdf"
    path.write_bytes(b"just some prose, no header")
    with pytest.raises(UnsupportedFormat):
        detect_format(path)


def test_html_saved_as_text_is_read_as_html(tmp_path):
    path = tmp_path / "filing.txt"
    path.write_text("<html><body><p>Revenue rose.</p></body></html>")
    assert detect_format(path) is DocumentFormat.HTML


def test_unknown_extension_is_rejected(tmp_path):
    path = tmp_path / "chart.png"
    path.write_bytes(b"\x89PNG\r\n\x1a\n")
    with pytest.raises(UnsupportedFormat):
        detect_format(path)


# -- Markdown table rendering --------------------------------------------- #

def test_render_table_drops_columns_that_are_empty_everywhere():
    """SEC filers pad statements with spacer cells; they carry no content."""
    rows = [["Line item", "", "2018", "", "2017"], ["Net income", "", "5,363", "", "4,869"]]
    rendered = render_table(rows)
    assert rendered.splitlines()[0] == "| Line item | 2018 | 2017 |"
    assert "5,363" in rendered


def test_render_table_escapes_pipes_so_a_cell_cannot_forge_a_column():
    rendered = render_table([["Segment", "Note"], ["Safety", "a | b"]])
    assert r"a \| b" in rendered


# -- CSV ------------------------------------------------------------------- #

def _write_csv(path, rows):
    with path.open("w", newline="") as handle:
        csv.writer(handle).writerows(rows)


def test_csv_becomes_one_table_segment_with_no_page_number(tmp_path):
    path = tmp_path / "revenue.csv"
    _write_csv(path, [["Segment", "FY2023"], ["Safety", "11020"], ["Transport", "8000"]])

    document = parse_document(path)

    assert document.source_format is DocumentFormat.CSV
    assert document.page_count == 1
    segment = document.pages[0]
    assert segment.segment_kind is SegmentKind.TABLE
    # The label names the file, not a page: a CSV has no page 1 to look up.
    assert "revenue" in segment.citation_label
    assert "| Safety | 11020 |" in segment.text
    assert document.is_paginated is False


def test_a_large_csv_is_cut_into_row_blocks_that_repeat_the_header(tmp_path):
    path = tmp_path / "ledger.csv"
    rows = [["Account", "Amount"]] + [[f"acct-{i}", str(i)] for i in range(600)]
    _write_csv(path, rows)

    document = parse_document(path)

    assert document.page_count == 3
    assert all(page.segment_kind is SegmentKind.SECTION for page in document.pages)
    assert "rows 2-201" in document.pages[0].citation_label
    assert "rows 402-601" in document.pages[2].citation_label
    # Without the header, a later block is an anonymous grid of numbers.
    assert "Account" in document.pages[2].text
    assert "acct-599" in document.pages[2].text


def test_tsv_is_split_on_tabs(tmp_path):
    path = tmp_path / "data.tsv"
    path.write_text("Segment\tFY2023\nSafety\t11020\n")
    document = parse_document(path)
    assert "| Safety | 11020 |" in document.pages[0].text


# -- Excel ----------------------------------------------------------------- #

def _workbook(path, sheets):
    from openpyxl import Workbook

    book = Workbook()
    book.remove(book.active)
    for title, rows in sheets.items():
        sheet = book.create_sheet(title)
        for row in rows:
            sheet.append(row)
    book.save(path)


def test_each_worksheet_is_its_own_segment(tmp_path):
    path = tmp_path / "segments.xlsx"
    _workbook(
        path,
        {
            "Q4 Revenue": [["Segment", "FY2023"], ["Safety", 11020]],
            "Notes": [["Note", "Detail"], [1, "Recognised at a point in time"]],
        },
    )

    document = parse_document(path)

    assert document.source_format is DocumentFormat.XLSX
    assert document.page_count == 2
    assert [page.segment_kind for page in document.pages] == [
        SegmentKind.SHEET,
        SegmentKind.SHEET,
    ]
    assert document.pages[0].citation_label == "sheet 'Q4 Revenue'"
    assert "11020" in document.pages[0].text
    assert document.pages[1].citation_label == "sheet 'Notes'"


def test_excel_formulas_are_read_as_their_cached_values(tmp_path):
    """An analyst asks what a cell says, not how it is computed."""
    from openpyxl import Workbook

    path = tmp_path / "calc.xlsx"
    book = Workbook()
    sheet = book.active
    sheet.title = "Totals"
    sheet.append(["Metric", "Value"])
    sheet.append(["Revenue", "=SUM(1,2)"])
    book.save(path)

    document = parse_document(path)
    # No cached value exists for a formula never opened in Excel, so the cell
    # reads empty rather than leaking "=SUM(1,2)" into the evidence.
    assert "=SUM" not in document.pages[0].text


# -- Word ------------------------------------------------------------------ #

def test_word_splits_on_explicit_page_breaks_only(tmp_path):
    from docx import Document as Docx

    path = tmp_path / "review.docx"
    document = Docx()
    document.add_heading("Annual Review", 1)
    document.add_paragraph("Revenue grew 4% year over year.")
    document.add_page_break()
    document.add_heading("Risk Factors", 1)
    document.add_paragraph("Supply chain disruption remains a risk.")
    document.save(path)

    parsed = parse_document(path)

    assert parsed.segmentation == "word-page-break"
    assert parsed.page_count == 2
    assert parsed.pages[0].segment_kind is SegmentKind.PAGE
    assert "# Annual Review" in parsed.pages[0].text
    assert "Risk Factors" in parsed.pages[1].text


def test_word_without_page_breaks_yields_sections_not_invented_pages(tmp_path):
    """No author break means no real page boundary, so nothing claims to be one."""
    from docx import Document as Docx

    path = tmp_path / "memo.docx"
    document = Docx()
    document.add_heading("Overview", 1)
    document.add_paragraph("First.")
    document.add_heading("Detail", 1)
    document.add_paragraph("Second.")
    document.save(path)

    parsed = parse_document(path)

    assert parsed.segmentation == "heading"
    assert all(page.segment_kind is SegmentKind.SECTION for page in parsed.pages)
    assert parsed.is_paginated is False


def test_word_tables_survive_as_markdown_tables(tmp_path):
    from docx import Document as Docx

    path = tmp_path / "table.docx"
    document = Docx()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "FY2023"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "$32,681M"
    document.save(path)

    parsed = parse_document(path)
    assert "| Revenue | $32,681M |" in parsed.pages[0].text


# -- HTML ------------------------------------------------------------------ #

def test_html_financial_tables_keep_figures_with_their_year_columns(tmp_path):
    path = tmp_path / "filing.htm"
    path.write_text(
        """
        <html><body>
        <table>
          <tr><td>(Millions)</td><td>2018</td><td>2017</td></tr>
          <tr><td>Purchases of PP&amp;E</td><td>(1,577)</td><td>(1,373)</td></tr>
        </table>
        </body></html>
        """
    )
    document = parse_document(path)
    text = document.pages[0].text
    assert "| Purchases of PP&E | (1,577) | (1,373) |" in text


def test_html_layout_tables_are_unwrapped_rather_than_rendered(tmp_path):
    """A one-column wrapper is page furniture; rendering it buries the content."""
    path = tmp_path / "wrapped.htm"
    path.write_text(
        "<html><body><table><tr><td>"
        "<p>Net sales increased 3%.</p>"
        "</td></tr></table></body></html>"
    )
    document = parse_document(path)
    text = document.pages[0].text
    assert "Net sales increased 3%." in text
    assert "|" not in text


# -- Markdown page store --------------------------------------------------- #

def test_markdown_store_writes_one_file_per_segment_and_reads_it_back(tmp_path):
    source = tmp_path / "segments.xlsx"
    _workbook(source, {"Revenue": [["Segment", "FY2023"], ["Safety", 11020]]})
    document = parse_document(source, doc_name="segments")

    store = MarkdownPageStore(base_dir=tmp_path / "markdown")
    store.save(document)

    directory = store.document_dir("segments")
    assert (directory / "manifest.json").exists()
    # A worksheet is stored as a sheet, not as a page.
    assert (directory / "sheet-001.md").exists()

    manifest = store.load_manifest("segments")
    assert manifest.source_format is DocumentFormat.XLSX
    assert manifest.segment_count == 1
    assert manifest.is_current

    restored = store.load_pages("segments")
    assert [page.text for page in restored] == [page.text for page in document.pages]
    assert restored[0].segment_kind is SegmentKind.SHEET
    assert restored[0].citation_label == "sheet 'Revenue'"


def test_markdown_store_uses_page_filenames_for_paginated_sources(tmp_path):
    source = tmp_path / "filing.htm"
    source.write_text(
        '<html><body><p>One</p><hr style="page-break-after: always">'
        "<p>Two</p></body></html>"
    )
    document = parse_document(source, doc_name="filing")

    store = MarkdownPageStore(base_dir=tmp_path / "markdown")
    store.save(document)

    directory = store.document_dir("filing")
    assert (directory / "page-001.md").exists()
    assert (directory / "page-002.md").exists()
    assert store.load_markdown("filing", 1) == "Two"


def test_re_saving_a_shorter_parse_leaves_no_stale_pages_behind(tmp_path):
    """A stale page-004.md would be read as current by anyone browsing the dir."""
    long_source = tmp_path / "filing.htm"
    long_source.write_text(
        '<html><body><p>One</p><hr style="page-break-after: always">'
        '<p>Two</p><hr style="page-break-after: always"><p>Three</p></body></html>'
    )
    store = MarkdownPageStore(base_dir=tmp_path / "markdown")
    store.save(parse_document(long_source, doc_name="filing"))
    assert (store.document_dir("filing") / "page-003.md").exists()

    short_source = tmp_path / "short.htm"
    short_source.write_text("<html><body><p>Only one page</p></body></html>")
    store.save(parse_document(short_source, doc_name="filing"))

    assert not (store.document_dir("filing") / "page-003.md").exists()
    assert store.load_manifest("filing").segment_count == 1
